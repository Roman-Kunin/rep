from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

document = Document()

document.add_heading("Заголовок_0", 0)
document.add_paragraph("I love Python")
style = document.styles.add_style("Style_1", WD_STYLE_TYPE.PARAGRAPH)
font = style.font
font.name("Times New Roman")


document.save(r"C:\Users\Roman\Desktop\Word\Python.doc")
